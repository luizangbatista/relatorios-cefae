import io
import os
import re
from datetime import date, datetime

import gspread
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt
from google.oauth2.service_account import Credentials
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import simpleSplit
from reportlab.pdfgen import canvas

# =========================================================
# CONFIGURAÇÃO INICIAL
# =========================================================

st.set_page_config(
    page_title="Relatórios de Monitoria",
    page_icon="📝",
    layout="wide",
)

SENHA_CORRETA = "*pazebem"
ARQUIVO_TIMBRADO = "timbrado.png"

# COLE AQUI O ID DA SUA PLANILHA
# Exemplo de URL:
# https://docs.google.com/spreadsheets/d/1m6zSOCCcV-Dz7BKBrEaHk77Vlq2AOyqLlh6adqDAXQY/edit?usp=drivesdk
ID_PLANILHA = "1m6zSOCCcV-Dz7BKBrEaHk77Vlq2AOyqLlh6adqDAXQY"

COLUNAS_ALUNOS = ["turma", "aluno"]
COLUNAS_RELATORIOS = ["data", "turma", "monitor", "alunos", "relatorio"]


MONITORES = [
    "Arthur - Matemática",
    "Davi - Ciências",
    "Dayane - História",
    "Gabriel - Física",
    "Gabriel - Português",
    "Lorraine - 4º ano",
    "Luiza - Matemática",
    "Maria Eduarda - 5º ano",
    "Rafael - Matemática",
    "Rayanne - 5º ano",
    "Roberta - 4º ano",
    "Silvana - Coordenação",
    "Uill - Português",
    "Vinícius - Inglês",
]

TURMAS_FIXAS = {
    "6º ano - Tarde": [
        "ALICE MORAIS DE FREITAS",
        "ANA ALICE RIBEIRO DE MELO",
        "ANA GABRIELLE G.DE CAMPOS RIBEIRO",
        "BERNARDO DE MELLO",
        "BERNARDO GABRIEL DA SILVA",
        "BRUNO BRASILEIRO COFFRAN",
        "DAVI FRANCISCO DOS SANTOS",
        "EMILLY CARDOSO COGUI",
        "GABRIEL FEITAS LEAL",
        "IYANLA GABRIELLE DIAS",
        "JOÃO VICENTE PIERUCCINI",
        "LUIZA MIRANDA ROSEMBRACK",
        "MABELLE PIETRA TEIXEIRA RIBEIRO",
        "MANUELA CARDOSO COTA",
        "MARIA LUISA DE SOUSA",
        "MATHEUS VASCONCELOS DIAS",
        "MIGUEL HOFFNER MOREIRA",
        "MIGUEL VITORI DIAS",
        "NATALIA NOGUEIRA SANTANA",
        "PAMELA GABRIELLE",
        "PEDRO QUEIROZ BESSA",
        "SAMUEL DOS SANTOS RAMOS",
        "THAYLA EMANUELLE RIBEIRO DA MATA",
        "VITORIA PINHEIRO SANTOS",
    ],
    "6º ano - Manhã": [
        "ANA LUIZA SOUSA LEMES SOARES",
        "DEBORA LORENA SOUSA LEMES SOARES",
        "GIOVANNA CORREA DE QUEIROZ",
        "HELENA CHAVES PIMENTA",
        "LARISSA ESTEVES PINHEIRO SANTOS",
        "LAURA TEIXEIRA BAUER",
        "LUIZA NEVES DOS SANTOS",
        "PEROLA GONÇALVES OLIVEIRA",
        "VITÓRIA EMANUELE BARBOSA DO ROSÁRIO",
    ],
    "5º ano - Manhã": [
        "ANA LUIZA EZEQUIEL LIMA",
        "HELENA FERNANDES VIEIRA",
        "JESSICA EMANUELLY G. DE OLIVEIRA",
        "JULIA BITENCOURT M. DOS ANJOS",
        "LAYLA DOS SANTOS",
        "NICOLAS GABRIEL ROCHA SOIER",
        "PEDRO HENRIQUE ALVES CRUZ",
        "GABRIELA VIEIRA CHAVES",
    ],
    "4º ano - Manhã": [
        "ANA VITORIA BATISTA DA SILVA",
        "EMANUELLY HELENA ALMEIDA LINO",
        "GEOVANNA SIQUEIRA DE SOUZA",
        "MARIA CLARA NEVES DORNELAS",
        "PIETRO MONTEIRO MARTINS",
        "SOPHIA EMANUELLY R. ESTEVES",
        "SOPHIA SOARES DANTAS",
        "YASMIN SARAH PEREIRA NUNES",
        "ISABELA DA CUNHA SILVA",
        "LAURA MARCELY VALENTIM",
    ],
    "4º ano - Tarde": [
        "ANA LUÍSA SANTOS DE JESUS",
        "ANA SOPHIA GARRIDO PÁSCOA",
        "APOLO PAIXAO NORONHA",
        "BRUNA ROMEIRO TOMÁS",
        "ENZO DE OLIVEIRA LIVRAMENTO",
        "GUILHERME VITORI DA COSTA",
        "KATARINA N. PEREIRA DA SILVA",
        "LUCAS GABRIEL OLIVEIRA",
        "MARIA EDUARDA FERREIRA SOUSA",
        "MARIA EDUARDA R. M. DE PAULA",
        "NUBIA GOMES TRINDADE",
        "REINALDO RIBERIO FILHO",
        "THEO ALONSO DE OLIVEIRA",
        "VALENTINA BRITO SILVA",
        "VITÓRIA E. MARTINS DOS SANTOS",
    ],
    "5º ano - Tarde": [
        "ALLAN GUIMARÃES",
        "ANTÔNIO CALDEIRA",
        "BEATRIZ ROMEIRO TOMAS",
        "EMANUELLE SANTOS CARIOCA MAGNANI",
        "GABRIEL N. ALVES DE MESQUITA",
        "ILANA AGATHA DIAS",
        "JOÃO VICTOR DE PAULA FELIX",
        "JULIA ASSIS PEREIRA",
        "LEONARDO CESAR NOVAIS LAGES",
        "LUCAS GABRIEL DA ROCHA",
        "MARIA LUISA DE SOUSA",
        "MIGUEL F. BARBOSA DE CARVALHO",
        "OTAVIO PIASSI MOTTA",
    ],
    "7º ano - Tarde": [
        "CLARICE SOARES DA CRUZ E SILVA",
        "ESTHER VIEIRA SANTOS ARAÚJO",
        "GABRIEL FRANCISCO SOUZA",
        "ÍGOR EMANUEL PEREIRA BRANDÃO",
        "ISAQUE SILVA REIS",
        "IZADORA SILVA DOS SANTOS",
        "LAVINIA CRISTINA SOUZA DA SILVA",
        "LEONARDO DO C. SANCHES ZAREMBA",
        "MIGUEL RIBEIRO DOS SANTOS",
        "SOPHIA RESENDE DE OLIVEIRA",
        "VITÓRIA SANTANA RIBEIRO LOMASSO",
    ],
    "8º ano - Tarde": [
        "ANA IZABEL RIBEIRO DE MELO",
        "DAVI DE MELO ALVES",
        "GABRIELLA SOUSA DAMAZIO",
        "HENRIQUE COELHO LACERDA",
        "ISABELLA VIANA DOS SANTOS",
        "JUAN CARLOS COSTA DE ALMEIDA",
        "LUCAS AGUIAR MARTINS",
        "LUIZA MARTINS COSTA MENDONÇA",
        "MARIA EDUARDA VALÉRIO COSTA",
    ],
    "9º ano - Tarde": [
        "ARTHUR CÉSAR VIEIRA",
        "BRUNA PAULA DE ALMIEDA",
        "CLARA ROCHA DE CARVALHO MARTINS",
        "GABRIEL JOSÉ BORONI DE MIRANDA",
        "GUSTAVO H. DAMASCENO PEREIRA",
        "IURI GABRIEL RODRIGUES MIRANDA",
        "JÚLIA GABRIELA CORDEIRO DE PAULA",
        "LAURA BEATRIZ SANTANA PEREIRA",
        "LAURA GUIMARÃES DA SILVA",
        "RÔMULO SILVEIRA DE SOUZA",
        "SOFIA FERREIRA BARBOSA",
        "SOPHIA VITORIA PEREIRA SANTOS",
        "YASMIN CALDEIRA COSTA",
    ],
    "1ª série - EM": [
        "ARTHUR FÉLIX RIBEIRO",
        "ARTHUR VIEIRA DE SOUZA BENTO",
        "BRENO HENRIQUE EUFRAZIO DE SOUZA",
        "CAUÃ TOMAGNINI JUNQUEIRA BARBOSA",
        "DANIEL HAMILTON COUTO CAMPOS",
        "DAVI LUCAS RAMOS DE ALMEIDA",
        "ENZO VIEIRA BAETA BATISTA MARTINS",
        "GIULIA TEIXEIRA ANDRADE",
        "HENRIQUE E. FERNANDES DOS SANTOS",
        "LETÍCIA DE ANDRADE LAGES",
        "LUAN MIGUEL VIEIRA DE SOUZA",
        "LUCAS PIERRE RIBEIRO PROCÓPIO",
        "LUISA DE OLIVEIRA SABARÁ",
        "LUIZA AGUIAR MARTINS",
        "MIGUEL KAIROS FELISBERTO AMORIM",
        "SARAH E. DE MORAIS SANTANA",
    ],
    "2ª série - EM": [
        "DAVI RAFAEL BORGES LOPES",
        "LUDMILA SANTANA LOMASSO",
        "MARIANA SANTOS",
        "SOLOS CASTRO CARMO NEVES",
        "STELLA MARTINS COSTA MENDONÇA",
    ],
}

if "autenticado" not in st.session_state:
    st.session_state.autenticado = False

if "tema" not in st.session_state:
    st.session_state.tema = "light"

if "pagina" not in st.session_state:
    st.session_state.pagina = "home"

if "mensagem_sucesso" not in st.session_state:
    st.session_state.mensagem_sucesso = ""

if "modo_exclusao" not in st.session_state:
    st.session_state.modo_exclusao = False

if "acesso_registrado" not in st.session_state:
    st.session_state.acesso_registrado = False


def alternar_tema():
    st.session_state.tema = "light" if st.session_state.tema == "dark" else "dark"


def obter_cores_tema():
    if st.session_state.tema == "dark":
        return {
            "BG": "#0b1220",
            "CARD": "#111827",
            "CARD_SOFT": "#0f172a",
            "BORDER": "#243041",
            "TEXT": "#f8fafc",
            "SUBTEXT": "#cbd5e1",
            "SUCCESS_BG": "#052e1a",
            "SUCCESS_BORDER": "#166534",
            "SUCCESS_TEXT": "#bbf7d0",
            "BUTTON_HOVER": "#172033",
        }

    return {
        "BG": "#f8fafc",
        "CARD": "#ffffff",
        "CARD_SOFT": "#f1f5f9",
        "BORDER": "#dbe4ee",
        "TEXT": "#111827",
        "SUBTEXT": "#4b5563",
        "SUCCESS_BG": "#ecfdf3",
        "SUCCESS_BORDER": "#86efac",
        "SUCCESS_TEXT": "#166534",
        "BUTTON_HOVER": "#eef2f7",
    }


def aplicar_estilo():
    cores = obter_cores_tema()
    st.markdown(
        f"""
        <style>
        .block-container {{
            padding-top: 1.1rem;
            padding-bottom: 2rem;
            max-width: 920px;
        }}
        html, body, [data-testid="stAppViewContainer"] {{
            background-color: {cores["BG"]};
        }}
        [data-testid="stHeader"] {{
            background: transparent;
        }}
        [data-testid="stToolbar"] {{
            right: 0.5rem;
        }}
        div[data-testid="stForm"] {{
            background: {cores["CARD"]};
            border: 1px solid {cores["BORDER"]};
            border-radius: 16px;
            padding: 1rem;
        }}
        .card-dark {{
            background: {cores["CARD"]};
            border: 1px solid {cores["BORDER"]};
            border-radius: 16px;
            padding: 1rem;
            margin-bottom: 1rem;
            box-shadow: 0 1px 0 rgba(0,0,0,0.02);
        }}
        .card-soft {{
            background: {cores["CARD_SOFT"]};
            border: 1px solid {cores["BORDER"]};
            border-radius: 16px;
            padding: 0.9rem;
            margin-bottom: 1rem;
        }}
        .home-title {{
            text-align: center;
            font-size: 2rem;
            font-weight: 700;
            margin-bottom: 0.35rem;
            color: {cores["TEXT"]};
        }}
        .home-subtitle {{
            text-align: center;
            font-size: 1rem;
            color: {cores["SUBTEXT"]};
            margin-bottom: 1.4rem;
        }}
        .success-box {{
            padding: 0.9rem 1rem;
            border-radius: 14px;
            background: {cores["SUCCESS_BG"]};
            border: 1px solid {cores["SUCCESS_BORDER"]};
            color: {cores["SUCCESS_TEXT"]};
            margin-bottom: 1rem;
            text-align: center;
            font-weight: 600;
        }}
        .section-title {{
            font-size: 1.05rem;
            font-weight: 700;
            color: {cores["TEXT"]};
            margin-bottom: 0.6rem;
        }}
        .status-box {{
            padding: 0.8rem 1rem;
            border-radius: 12px;
            background: {cores["CARD_SOFT"]};
            border: 1px solid {cores["BORDER"]};
            margin-top: 1rem;
            text-align: center;
        }}
        .stButton > button,
        .stDownloadButton > button {{
            width: 100%;
            border-radius: 14px;
            min-height: 52px;
            font-size: 16px;
            background: {cores["CARD"]};
            color: {cores["TEXT"]};
            border: 1px solid {cores["BORDER"]};
        }}
        .stButton > button:hover,
        .stDownloadButton > button:hover {{
            border-color: #60a5fa;
            color: {cores["TEXT"]};
            background: {cores["BUTTON_HOVER"]};
        }}
        div[data-baseweb="select"] > div,
        div[data-baseweb="input"] > div,
        div[data-baseweb="textarea"] > div {{
            background-color: {cores["CARD"]} !important;
            border-color: {cores["BORDER"]} !important;
            color: {cores["TEXT"]} !important;
            border-radius: 12px !important;
        }}
        input, textarea {{
            color: {cores["TEXT"]} !important;
        }}
        label, .stMarkdown, .stText, p, span, div {{
            color: {cores["TEXT"]};
        }}
        div[data-testid="stDateInput"] > div {{
            background-color: {cores["CARD"]} !important;
            border-radius: 12px !important;
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )


aplicar_estilo()

def dataframe_alunos_fixo():
    linhas = []
    for turma, alunos in TURMAS_FIXAS.items():
        for aluno in alunos:
            linhas.append({"turma": turma, "aluno": aluno})
    return pd.DataFrame(linhas, columns=COLUNAS_ALUNOS)


def obter_id_planilha(valor):
    valor = str(valor or "").strip()

    if not valor or valor == "COLE_AQUI_O_ID_DA_PLANILHA":
        raise ValueError("Preencha o ID_PLANILHA com o ID ou a URL da sua planilha do Google Sheets.")

    correspondencia = re.search(r"/spreadsheets/d/([a-zA-Z0-9-_]+)", valor)
    if correspondencia:
        return correspondencia.group(1)

    if re.fullmatch(r"[a-zA-Z0-9-_]{20,}", valor):
        return valor

    raise ValueError("ID_PLANILHA inválido. Use apenas o ID da planilha ou a URL completa do Google Sheets.")


def obter_credenciais_google():
    if "gcp_service_account" not in st.secrets:
        raise KeyError(
            "Credenciais do Google não encontradas no Streamlit Secrets. "
            "Adicione a seção [gcp_service_account] no arquivo secrets.toml."
        )

    info = dict(st.secrets["gcp_service_account"])

    chave_privada = str(info.get("private_key", "")).strip()
    if not chave_privada or "sua chave aqui" in chave_privada.lower():
        raise ValueError(
            "A private_key do service account está vazia ou é apenas um placeholder. "
            "Cole a chave JSON real no secrets.toml."
        )

    info["private_key"] = chave_privada.replace("\\n", "\n")

    scopes = [
        "https://www.googleapis.com/auth/spreadsheets",
        "https://www.googleapis.com/auth/drive",
    ]
    return Credentials.from_service_account_info(info, scopes=scopes)


def obter_ou_criar_aba(planilha, nome_aba, colunas):
    try:
        aba = planilha.worksheet(nome_aba)
    except gspread.WorksheetNotFound:
        aba = planilha.add_worksheet(title=nome_aba, rows=max(100, len(colunas) + 5), cols=max(10, len(colunas) + 2))

    valores = aba.get_all_values()
    cabecalho_atual = valores[0] if valores else []

    if cabecalho_atual != colunas:
        conteudo_existente = valores[1:] if valores else []
        aba.clear()
        aba.update("A1", [colunas] + conteudo_existente)

    return aba


def conectar_google_sheets():
    creds = obter_credenciais_google()
    client = gspread.authorize(creds)
    planilha = client.open_by_key(obter_id_planilha(ID_PLANILHA))

    obter_ou_criar_aba(planilha, "relatorios", COLUNAS_RELATORIOS)

    return planilha

def carregar_relatorios():
    planilha = conectar_google_sheets()
    aba = planilha.worksheet("relatorios")
    dados = aba.get_all_records()
    df = pd.DataFrame(dados)

    if df.empty:
        df = pd.DataFrame(columns=COLUNAS_RELATORIOS)
        df["data_dt"] = pd.to_datetime(pd.Series(dtype="object"))
        return df

    for col in COLUNAS_RELATORIOS:
        if col not in df.columns:
            df[col] = ""
        df[col] = df[col].fillna("").astype(str).str.strip()

    df["data_dt"] = pd.to_datetime(df["data"], errors="coerce")
    return df.sort_values("data_dt", ascending=False).reset_index(drop=True)



def carregar_alunos():
    return dataframe_alunos_fixo()

def salvar_relatorio(data_relatorio, turma, monitor, alunos, texto_relatorio):
    turma = str(turma).strip()
    monitor = str(monitor).strip()
    texto_relatorio = str(texto_relatorio).strip()
    alunos_texto = "; ".join(alunos)

    if not turma:
        return False, "Selecione uma turma."
    if not monitor:
        return False, "Selecione um monitor."
    if not alunos:
        return False, "Selecione pelo menos um aluno."
    if not texto_relatorio:
        return False, "Escreva o relatório."

    try:
        planilha = conectar_google_sheets()
        aba = planilha.worksheet("relatorios")
        aba.append_row(
            [
                data_relatorio.strftime("%Y-%m-%d"),
                turma,
                monitor,
                alunos_texto,
                texto_relatorio,
            ],
            value_input_option="USER_ENTERED",
        )
        return True, "Relatório salvo com sucesso."
    except Exception as e:
        return False, f"Erro ao salvar: {str(e)}"


def filtrar_relatorios(df, turma=None, aluno=None, monitor=None, data_ini=None, data_fim=None):
    if df.empty:
        return df.copy()

    filtrado = df.copy()

    if turma and turma != "Todas":
        filtrado = filtrado[filtrado["turma"] == turma]

    if monitor and monitor != "Todos":
        filtrado = filtrado[filtrado["monitor"] == monitor]

    if aluno and aluno != "Todos":
        filtrado = filtrado[
            filtrado["alunos"].apply(
                lambda x: aluno in [parte.strip() for parte in str(x).split(";") if parte.strip()]
            )
        ]

    if data_ini:
        filtrado = filtrado[filtrado["data_dt"] >= pd.Timestamp(data_ini)]

    if data_fim:
        filtrado = filtrado[filtrado["data_dt"] <= pd.Timestamp(data_fim)]

    return filtrado.sort_values("data_dt", ascending=False).reset_index(drop=True)


def gerar_texto_filtros_utilizados(turma_filtro, aluno_filtro, monitor_filtro, data_ini, data_fim):
    filtros = []

    if turma_filtro and turma_filtro != "Todas":
        filtros.append(f"Turma: {turma_filtro}")

    if aluno_filtro and aluno_filtro != "Todos":
        filtros.append(f"Aluno: {aluno_filtro}")

    if monitor_filtro and monitor_filtro != "Todos":
        filtros.append(f"Monitor: {monitor_filtro}")

    if data_ini:
        filtros.append(f"Data inicial: {data_ini.strftime('%d/%m/%Y')}")

    if data_fim:
        filtros.append(f"Data final: {data_fim.strftime('%d/%m/%Y')}")

    return " | ".join(filtros) if filtros else ""


def deletar_relatorios(df_filtrado, indices_filtrados):
    if not indices_filtrados:
        return False, "Selecione pelo menos um relatório para excluir."

    try:
        planilha = conectar_google_sheets()
        aba = planilha.worksheet("relatorios")
        valores = aba.get_all_values()

        if not valores or len(valores) < 2:
            return False, "Nenhum relatório encontrado para excluir."

        cabecalho = valores[0]
        linhas = valores[1:]
        df_completo = pd.DataFrame(linhas, columns=cabecalho)

        for col in COLUNAS_RELATORIOS:
            if col not in df_completo.columns:
                df_completo[col] = ""

        df_completo = df_completo[COLUNAS_RELATORIOS].copy()
        linhas_para_remover = df_filtrado.loc[indices_filtrados, COLUNAS_RELATORIOS].copy()
        restantes = df_completo.copy()

        for _, linha in linhas_para_remover.iterrows():
            mascara = (
                (restantes["data"] == str(linha["data"])) &
                (restantes["turma"] == str(linha["turma"])) &
                (restantes["monitor"] == str(linha["monitor"])) &
                (restantes["alunos"] == str(linha["alunos"])) &
                (restantes["relatorio"] == str(linha["relatorio"]))
            )
            idx_match = restantes[mascara].index
            if len(idx_match) > 0:
                restantes = restantes.drop(idx_match[0])

        aba.clear()
        aba.append_row(COLUNAS_RELATORIOS)

        if not restantes.empty:
            aba.append_rows(
                restantes[COLUNAS_RELATORIOS].values.tolist(),
                value_input_option="USER_ENTERED",
            )

        return True, f"{len(indices_filtrados)} relatório(s) excluído(s) com sucesso."
    except Exception as e:
        return False, f"Erro ao excluir: {str(e)}"


def gerar_pdf_relatorios(df, filtros_texto):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)

    largura, altura = A4
    margem_esq = 1.5 * 28.35
    margem_dir = 1.5 * 28.35
    margem_topo = 4.5 * 28.35
    margem_base = 4.0 * 28.35
    largura_texto = largura - margem_esq - margem_dir

    fonte_normal = "Helvetica"
    fonte_negrito = "Helvetica-Bold"
    tamanho = 11
    espacamento_linha = 16.5
    espacamento_relatorio = 28.35

    y = altura - margem_topo

    def desenhar_timbrado():
        if os.path.exists(ARQUIVO_TIMBRADO):
            c.drawImage(ARQUIVO_TIMBRADO, 0, 0, width=largura, height=altura)

    def nova_pagina():
        nonlocal y
        c.showPage()
        desenhar_timbrado()
        y = altura - margem_topo

    def escrever_linha_centralizada(texto, fonte=fonte_normal, tamanho_fonte=tamanho, espaco=18):
        nonlocal y
        linhas = simpleSplit(str(texto), fonte, tamanho_fonte, largura_texto)
        c.setFont(fonte, tamanho_fonte)

        for linha in linhas:
            if y < margem_base:
                nova_pagina()
                c.setFont(fonte, tamanho_fonte)

            largura_linha = c.stringWidth(linha, fonte, tamanho_fonte)
            x = margem_esq + (largura_texto - largura_linha) / 2
            c.drawString(x, y, linha)
            y -= espaco

    def escrever_texto_justificado(rotulo, texto, espaco_linha=espacamento_linha):
        nonlocal y

        rotulo = str(rotulo)
        texto = str(texto).strip()

        c.setFont(fonte_negrito, tamanho)
        largura_rotulo = c.stringWidth(rotulo, fonte_negrito, tamanho)
        largura_primeira = largura_texto - largura_rotulo

        palavras = texto.split()
        linhas = []
        linha_atual = ""
        largura_limite = largura_primeira

        c.setFont(fonte_normal, tamanho)

        for palavra in palavras:
            teste = palavra if not linha_atual else f"{linha_atual} {palavra}"
            if c.stringWidth(teste, fonte_normal, tamanho) <= largura_limite:
                linha_atual = teste
            else:
                if linha_atual:
                    linhas.append(linha_atual)
                linha_atual = palavra
                largura_limite = largura_texto

        if linha_atual:
            linhas.append(linha_atual)

        if not linhas:
            linhas = [""]

        for i, linha in enumerate(linhas):
            if y < margem_base:
                nova_pagina()

            if i == 0:
                c.setFont(fonte_negrito, tamanho)
                c.drawString(margem_esq, y, rotulo)
                c.setFont(fonte_normal, tamanho)
                c.drawString(margem_esq + largura_rotulo, y, linha)
            else:
                palavras_linha = linha.split()

                if len(palavras_linha) > 1 and i != len(linhas) - 1:
                    largura_palavras = sum(c.stringWidth(p, fonte_normal, tamanho) for p in palavras_linha)
                    espaco_total = largura_texto - largura_palavras
                    espaco = espaco_total / (len(palavras_linha) - 1)

                    x = margem_esq
                    c.setFont(fonte_normal, tamanho)

                    for palavra in palavras_linha[:-1]:
                        c.drawString(x, y, palavra)
                        x += c.stringWidth(palavra, fonte_normal, tamanho) + espaco

                    c.drawString(x, y, palavras_linha[-1])
                else:
                    c.setFont(fonte_normal, tamanho)
                    c.drawString(margem_esq, y, linha)

            y -= espaco_linha

    def escrever_linha_mista(data_str, monitor, turma, espaco=espacamento_linha):
        nonlocal y

        partes = [
            ("Data:", True),
            (f" {data_str} | ", False),
            ("Monitor:", True),
            (f" {monitor} | ", False),
            ("Turma:", True),
            (f" {turma}", False),
        ]

        linhas = []
        linha_atual = []
        largura_atual = 0

        for texto, negrito in partes:
            fonte = fonte_negrito if negrito else fonte_normal
            largura_parte = c.stringWidth(texto, fonte, tamanho)

            if largura_atual + largura_parte <= largura_texto:
                linha_atual.append((texto, negrito))
                largura_atual += largura_parte
            else:
                if linha_atual:
                    linhas.append(linha_atual)
                linha_atual = [(texto, negrito)]
                largura_atual = largura_parte

        if linha_atual:
            linhas.append(linha_atual)

        for linha in linhas:
            if y < margem_base:
                nova_pagina()

            x = margem_esq
            for texto, negrito in linha:
                fonte = fonte_negrito if negrito else fonte_normal
                c.setFont(fonte, tamanho)
                c.drawString(x, y, texto)
                x += c.stringWidth(texto, fonte, tamanho)

            y -= espaco

    desenhar_timbrado()

    if filtros_texto:
        escrever_linha_centralizada(filtros_texto, fonte_normal, tamanho, 18)
        y -= 10

    if df.empty:
        c.setFont(fonte_negrito, tamanho)
        c.drawString(margem_esq, y, "Nenhum relatório encontrado.")
    else:
        for i, (_, row) in enumerate(df.iterrows()):
            try:
                data_formatada = pd.to_datetime(row["data"]).strftime("%d/%m")
            except Exception:
                data_formatada = str(row["data"])

            if i > 0:
                y -= espacamento_relatorio

            escrever_linha_mista(data_formatada, str(row.get("monitor", "")), str(row.get("turma", "")))
            escrever_texto_justificado("Alunos: ", str(row.get("alunos", "")))
            escrever_texto_justificado("Relatório: ", str(row.get("relatorio", "")))

    c.save()
    buffer.seek(0)
    return buffer


def gerar_docx_relatorios(df, filtros_texto):
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Pt(42.52)
    sec.bottom_margin = Pt(42.52)
    sec.left_margin = Pt(42.52)
    sec.right_margin = Pt(42.52)

    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Calibri"
    estilo_normal.font.size = Pt(11)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = 1
    p_titulo.paragraph_format.space_after = Pt(6)

    r_titulo = p_titulo.add_run("RELATÓRIOS CEFAE - Arquivo editável")
    r_titulo.bold = True
    r_titulo.font.name = "Calibri"
    r_titulo.font.size = Pt(14)

    if filtros_texto:
        p_filtros = doc.add_paragraph()
        p_filtros.alignment = 1
        p_filtros.paragraph_format.line_spacing = 1.5
        p_filtros.paragraph_format.space_after = Pt(14)

        r_filtros = p_filtros.add_run(filtros_texto)
        r_filtros.font.name = "Calibri"
        r_filtros.font.size = Pt(11)
        r_filtros.italic = True
    else:
        p_espaco = doc.add_paragraph()
        p_espaco.paragraph_format.space_after = Pt(14)

    if df.empty:
        p = doc.add_paragraph()
        r = p.add_run("Nenhum relatório encontrado.")
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(11)
    else:
        for i, (_, row) in enumerate(df.iterrows()):
            try:
                data_convertida = pd.to_datetime(row.get("data", ""), errors="coerce")
                data_formatada = data_convertida.strftime("%d/%m") if pd.notna(data_convertida) else str(row.get("data", ""))
            except Exception:
                data_formatada = str(row.get("data", ""))

            if i > 0:
                p_espaco = doc.add_paragraph()
                p_espaco.paragraph_format.space_before = Pt(28.35)

            p1 = doc.add_paragraph()
            p1.paragraph_format.line_spacing = 1.5

            r = p1.add_run("Data: ")
            r.bold = True
            r.font.name = "Calibri"
            r.font.size = Pt(11)

            r = p1.add_run(f"{data_formatada} | ")
            r.font.name = "Calibri"
            r.font.size = Pt(11)

            r = p1.add_run("Monitor: ")
            r.bold = True
            r.font.name = "Calibri"
            r.font.size = Pt(11)

            r = p1.add_run(f"{row.get('monitor', '')} | ")
            r.font.name = "Calibri"
            r.font.size = Pt(11)

            r = p1.add_run("Turma: ")
            r.bold = True
            r.font.name = "Calibri"
            r.font.size = Pt(11)

            r = p1.add_run(str(row.get("turma", "")))
            r.font.name = "Calibri"
            r.font.size = Pt(11)

            p2 = doc.add_paragraph()
            p2.paragraph_format.line_spacing = 1.5

            r = p2.add_run("Alunos: ")
            r.bold = True
            r.font.name = "Calibri"
            r.font.size = Pt(11)

            r = p2.add_run(str(row.get("alunos", "")))
            r.font.name = "Calibri"
            r.font.size = Pt(11)

            p3 = doc.add_paragraph()
            p3.paragraph_format.line_spacing = 1.5
            p3.paragraph_format.alignment = 3

            r = p3.add_run("Relatório: ")
            r.bold = True
            r.font.name = "Calibri"
            r.font.size = Pt(11)

            r = p3.add_run(str(row.get("relatorio", "")))
            r.font.name = "Calibri"
            r.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def ir_para(nome_pagina):
    st.session_state.pagina = nome_pagina
    st.rerun()


def sair():
    st.session_state.autenticado = False
    st.session_state.pagina = "home"
    st.session_state.modo_exclusao = False
    st.rerun()


def topo_app():
    c1, c2, c3 = st.columns([7, 1, 1])

    with c2:
        st.button(
            "🌙" if st.session_state.tema == "dark" else "☀️",
            on_click=alternar_tema,
            use_container_width=True,
        )

    with c3:
        st.button("Sair", on_click=sair, use_container_width=True)


def botao_voltar():
    if st.button("⬅️ Voltar para a página inicial"):
        st.session_state.modo_exclusao = False
        ir_para("home")


def tela_login():
    _, topo2 = st.columns([8, 1])

    with topo2:
        st.button(
            "🌙" if st.session_state.tema == "dark" else "☀️",
            on_click=alternar_tema,
            use_container_width=True,
        )

    st.markdown('<div class="card-dark">', unsafe_allow_html=True)
    st.markdown('<div class="home-title">🔒 Acesso restrito</div>', unsafe_allow_html=True)
    st.markdown('<div class="home-subtitle">Digite a senha para acessar o sistema</div>', unsafe_allow_html=True)

    senha = st.text_input("Senha", type="password")

    if st.button("Entrar", use_container_width=True):
        if senha == SENHA_CORRETA:
            st.session_state.autenticado = True
            st.rerun()
        else:
            st.error("Senha incorreta")

    st.markdown("</div>", unsafe_allow_html=True)


def tela_home():
    topo_app()

    st.markdown('<div class="home-title">📚 Sistema de Monitoria</div>', unsafe_allow_html=True)
    st.markdown('<div class="home-subtitle">Selecione uma das opções abaixo</div>', unsafe_allow_html=True)

    if st.session_state.mensagem_sucesso:
        st.markdown(f'<div class="success-box">{st.session_state.mensagem_sucesso}</div>', unsafe_allow_html=True)
        st.session_state.mensagem_sucesso = ""

    c1, c2 = st.columns(2)

    with c1:
        if st.button("Enviar novo relatório", use_container_width=True):
            st.session_state.modo_exclusao = False
            ir_para("cadastrar_relatorio")

    with c2:
        if st.button("Consultar relatórios enviados", use_container_width=True):
            st.session_state.modo_exclusao = False
            ir_para("consultar")




def tela_cadastrar_relatorio():
    topo_app()
    botao_voltar()
    st.title("Enviar novo relatório")

    turmas_disponiveis = list(TURMAS_FIXAS.keys())

    if not turmas_disponiveis:
        st.warning("Nenhuma turma disponível.")
        return

    st.markdown('<div class="card-dark">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Cadastro de relatório</div>', unsafe_allow_html=True)

    col1, col2, col3 = st.columns(3)

    with col1:
        data_relatorio = st.date_input("Data", value=date.today(), format="DD/MM/YYYY")

    with col2:
        turma_escolhida = st.selectbox("Turma", options=turmas_disponiveis)

    with col3:
        monitor_escolhido = st.selectbox("Monitor", options=MONITORES)

    alunos_da_turma = sorted(TURMAS_FIXAS.get(turma_escolhida, []))

    st.markdown("### Seleção de alunos")

    chave_alunos = f"alunos_selecionados_{turma_escolhida}"
    if chave_alunos not in st.session_state:
        st.session_state[chave_alunos] = []

    b1, b2 = st.columns(2)

    with b1:
        if st.button("Selecionar todos"):
            st.session_state[chave_alunos] = alunos_da_turma.copy()
            st.rerun()

    with b2:
        if st.button("Selecionar nenhum"):
            st.session_state[chave_alunos] = []
            st.rerun()

    alunos_selecionados = st.multiselect(
        "Alunos da turma",
        options=alunos_da_turma,
        default=st.session_state[chave_alunos],
    )
    st.session_state[chave_alunos] = alunos_selecionados

    texto_relatorio = st.text_area("Relatório", height=220, placeholder="Escreva aqui o relatório da monitoria...")

    if st.button("Salvar relatório"):
        ok, mensagem = salvar_relatorio(
            data_relatorio=data_relatorio,
            turma=turma_escolhida,
            monitor=monitor_escolhido,
            alunos=alunos_selecionados,
            texto_relatorio=texto_relatorio,
        )

        if ok:
            st.session_state[chave_alunos] = []
            st.session_state.mensagem_sucesso = mensagem
            st.session_state.modo_exclusao = False
            ir_para("home")
        else:
            st.error(mensagem)

    st.markdown("</div>", unsafe_allow_html=True)


def tela_consultar(df_relatorios):
    topo_app()
    botao_voltar()
    st.title("Consultar relatórios enviados")

    if df_relatorios.empty:
        st.info("Nenhum relatório cadastrado ainda.")
        return

    st.markdown('<div class="card-dark">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Filtros</div>', unsafe_allow_html=True)

    turmas = ["Todas"] + sorted(df_relatorios["turma"].dropna().astype(str).unique().tolist())
    monitores = ["Todos"] + sorted(df_relatorios["monitor"].dropna().astype(str).unique().tolist())

    todos_alunos = set()
    for texto in df_relatorios["alunos"].fillna("").astype(str):
        for parte in texto.split(";"):
            nome = parte.strip()
            if nome:
                todos_alunos.add(nome)

    alunos_filtro = ["Todos"] + sorted(todos_alunos)

    c1, c2, c3 = st.columns(3)

    with c1:
        turma_filtro = st.selectbox("Filtrar por turma", options=turmas)

    with c2:
        aluno_filtro = st.selectbox("Filtrar por aluno", options=alunos_filtro)

    with c3:
        monitor_filtro = st.selectbox("Filtrar por monitor", options=monitores)

    usar_filtro_data = st.checkbox("Filtrar por data")

    if usar_filtro_data:
        c4, c5 = st.columns(2)
        with c4:
            data_ini = st.date_input("Data inicial", value=date.today(), format="DD/MM/YYYY", key="data_ini_consulta")
        with c5:
            data_fim = st.date_input("Data final", value=date.today(), format="DD/MM/YYYY", key="data_fim_consulta")
    else:
        data_ini = None
        data_fim = None

    st.markdown("</div>", unsafe_allow_html=True)

    df_filtrado = filtrar_relatorios(
        df=df_relatorios,
        turma=turma_filtro,
        aluno=aluno_filtro,
        monitor=monitor_filtro,
        data_ini=data_ini,
        data_fim=data_fim,
    )

    filtros_texto = gerar_texto_filtros_utilizados(turma_filtro, aluno_filtro, monitor_filtro, data_ini, data_fim)

    st.markdown('<div class="card-dark">', unsafe_allow_html=True)
    st.markdown(f"**Total encontrado:** {len(df_filtrado)} relatório(s)")

    if df_filtrado.empty:
        st.warning("Nenhum relatório corresponde aos filtros selecionados.")
        st.session_state.modo_exclusao = False
        st.markdown("</div>", unsafe_allow_html=True)
        return

    c_download1, c_download2 = st.columns(2)

    with c_download1:
        pdf_bytes = gerar_pdf_relatorios(df_filtrado, filtros_texto)
        st.download_button(
            label="Gerar arquivo PDF",
            data=pdf_bytes,
            file_name=f"relatorio_cefae_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )

    with c_download2:
        docx_bytes = gerar_docx_relatorios(df_filtrado, filtros_texto)
        st.download_button(
            label="Gerar arquivo DOCX",
            data=docx_bytes,
            file_name=f"relatorio_cefae_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    c_botao1, c_botao2 = st.columns(2)

    with c_botao1:
        if not st.session_state.modo_exclusao:
            if st.button("Excluir relatórios", use_container_width=True):
                st.session_state.modo_exclusao = True
                st.rerun()

    with c_botao2:
        if st.session_state.modo_exclusao:
            if st.button("Cancelar exclusão", use_container_width=True):
                st.session_state.modo_exclusao = False
                st.rerun()

    st.markdown("</div>", unsafe_allow_html=True)

    if st.session_state.modo_exclusao:
        st.markdown('<div class="card-soft">', unsafe_allow_html=True)
        st.markdown("### Selecione os relatórios para excluir")
        st.markdown("</div>", unsafe_allow_html=True)

    indices_para_excluir = []

    for idx, row in df_filtrado.iterrows():
        try:
            data_formatada = pd.to_datetime(row["data"]).strftime("%d/%m/%Y")
        except Exception:
            data_formatada = str(row["data"])

        st.markdown('<div class="card-dark">', unsafe_allow_html=True)

        if st.session_state.modo_exclusao:
            marcado = st.checkbox("Selecionar para excluir", key=f"excluir_relatorio_{idx}")
            if marcado:
                indices_para_excluir.append(idx)

        st.write(f"**Data:** {data_formatada}")
        st.write(f"**Turma:** {row['turma']}")
        st.write(f"**Monitor:** {row['monitor']}")
        st.write(f"**Alunos:** {row['alunos']}")
        st.write(f"**Relatório:** {row['relatorio']}")

        st.markdown("</div>", unsafe_allow_html=True)

    if st.session_state.modo_exclusao:
        if st.button("Excluir selecionados", use_container_width=True):
            ok, mensagem = deletar_relatorios(df_filtrado, indices_para_excluir)

            if ok:
                for idx in df_filtrado.index:
                    chave = f"excluir_relatorio_{idx}"
                    if chave in st.session_state:
                        del st.session_state[chave]

                st.session_state.modo_exclusao = False
                st.session_state.mensagem_sucesso = mensagem
                ir_para("consultar")
            else:
                st.error(mensagem)


if not st.session_state.autenticado:
    tela_login()
    st.stop()

df_alunos = carregar_alunos()
df_relatorios = carregar_relatorios()
pagina = st.session_state.pagina

if pagina == "home":
    tela_home()
elif pagina == "cadastrar_relatorio":
    tela_cadastrar_relatorio()
elif pagina == "consultar":
    tela_consultar(df_relatorios)
